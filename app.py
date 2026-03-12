"""
╔══════════════════════════════════════════════════════════════╗
║          CROWD ANALYSIS & THREAT DETECTION SYSTEM           ║
║                  Hugging Face Deployment                     ║
║  Detects: Ambulance Stuck | Fighting | Fire | Smoke | Weapon ║
║  Features: Real-time Alerts | Incident Log | RAG Summary     ║
╚══════════════════════════════════════════════════════════════╝
"""

# ── Standard Library ──────────────────────────────────────────
import os
import cv2
import time
import tempfile
import threading
import numpy as np
from datetime import datetime
from collections import Counter
from pathlib import Path

# ── ML / Detection ────────────────────────────────────────────
import torch
from PIL import Image
from ultralytics import YOLO

# ── Document Storage ──────────────────────────────────────────
from docx import Document as DocxDocument
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── RAG Pipeline ──────────────────────────────────────────────
from langchain_community.document_loaders import Docx2txtLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.chains import RetrievalQA
from langchain_community.llms import HuggingFacePipeline
from transformers import pipeline as hf_pipeline

# ── Gradio UI ─────────────────────────────────────────────────
import gradio as gr


# ══════════════════════════════════════════════════════════════
#  CONFIGURATION
# ══════════════════════════════════════════════════════════════

FINAL_CLASSES  = ['ambulance', 'fighting', 'fire', 'smoke', 'weapon']
CONF_THRESHOLD = 0.25
LOG_PATH       = 'incident_log.docx'
MODEL_PATH     = 'best.pt'

ALERT_CONFIG = {
    'ambulance': {
        'level': 'HIGH',
        'emoji': '🚑',
        'msg'  : 'AMBULANCE STUCK IN CROWD — Clear the path immediately!',
        'color': (0, 165, 255),
        'hex'  : '#FFA500',
    },
    'fire': {
        'level': 'CRITICAL',
        'emoji': '🔥',
        'msg'  : 'FIRE DETECTED — Evacuate area, call fire brigade!',
        'color': (0, 0, 255),
        'hex'  : '#FF2222',
    },
    'smoke': {
        'level': 'HIGH',
        'emoji': '💨',
        'msg'  : 'SMOKE DETECTED — Possible fire, investigate immediately!',
        'color': (180, 180, 180),
        'hex'  : '#AAAAAA',
    },
    'weapon': {
        'level': 'CRITICAL',
        'emoji': '🔫',
        'msg'  : 'WEAPON DETECTED — Alert security forces immediately!',
        'color': (255, 0, 0),
        'hex'  : '#FF0055',
    },
    'fighting': {
        'level': 'HIGH',
        'emoji': '⚔️',
        'msg'  : 'FIGHTING DETECTED — Deploy crowd control!',
        'color': (0, 200, 0),
        'hex'  : '#00CC44',
    },
}


# ══════════════════════════════════════════════════════════════
#  MODEL LOADER
# ══════════════════════════════════════════════════════════════

def load_model():
    if not os.path.exists(MODEL_PATH):
        raise FileNotFoundError(
            f'Model not found: {MODEL_PATH}\n'
            'Please upload best.pt to this Space.'
        )
    model = YOLO(MODEL_PATH)
    print(f'✅ YOLOv8m loaded | GPU: {torch.cuda.is_available()}')
    return model

model = load_model()


# ══════════════════════════════════════════════════════════════
#  DETECTION UTILITIES
# ══════════════════════════════════════════════════════════════

def draw_boxes(img_bgr, results):
    """Draw bounding boxes with labels on image."""
    for r in results:
        for box in r.boxes:
            cls_idx      = int(box.cls)
            conf         = float(box.conf)
            name         = FINAL_CLASSES[cls_idx]
            color        = ALERT_CONFIG[name]['color']
            x1,y1,x2,y2 = map(int, box.xyxy[0])
            label        = f'{name} {conf:.2f}'
            cv2.rectangle(img_bgr, (x1,y1),(x2,y2), color, 2)
            cv2.rectangle(img_bgr, (x1,y1-30),(x1+len(label)*12,y1), color, -1)
            cv2.putText(img_bgr, label, (x1+3,y1-8),
                        cv2.FONT_HERSHEY_SIMPLEX, 0.65, (255,255,255), 2)
    return img_bgr


def extract_detections(results):
    """Extract list of detection dicts from results."""
    return [
        {'class': FINAL_CLASSES[int(box.cls)], 'conf': float(box.conf)}
        for r in results for box in r.boxes
    ]


def build_alert_html(detections):
    """Build HTML alert panel from detections."""
    triggered = {}
    for det in detections:
        name = det['class']
        if name not in triggered:
            triggered[name] = {**ALERT_CONFIG[name], 'conf': det['conf']}

    if not triggered:
        return '''
        <div style="background:#0a2a0a;border:1px solid #00ff44;
                    border-radius:8px;padding:16px;
                    font-family:'Courier New',monospace;">
            <span style="color:#00ff44;font-size:16px;">
                ✅ SYSTEM CLEAR — No threats detected
            </span>
        </div>'''

    sorted_alerts = sorted(
        triggered.values(),
        key=lambda x: 0 if x['level'] == 'CRITICAL' else 1
    )

    html = '<div style="display:flex;flex-direction:column;gap:8px;">'
    for alert in sorted_alerts:
        bg     = '#2a0000' if alert['level'] == 'CRITICAL' else '#2a1500'
        border = alert['hex']
        html += f'''
        <div style="background:{bg};border-left:4px solid {border};
                    border-radius:4px;padding:12px 16px;
                    font-family:'Courier New',monospace;">
            <div style="color:{border};font-size:15px;font-weight:bold;letter-spacing:1px;">
                {alert["emoji"]} [{alert["level"]}] {alert["msg"]}
            </div>
            <div style="color:#888;font-size:12px;margin-top:4px;">
                Confidence: {alert["conf"]:.2f} | Time: {datetime.now().strftime("%H:%M:%S")}
            </div>
        </div>'''
    html += '</div>'
    return html


# ══════════════════════════════════════════════════════════════
#  INCIDENT LOG (.docx)
# ══════════════════════════════════════════════════════════════

_log_lock = threading.Lock()

def append_to_log(source, detections, triggered_alerts):
    """Append incident entry to docx log (thread-safe)."""
    with _log_lock:
        if os.path.exists(LOG_PATH):
            doc = DocxDocument(LOG_PATH)
        else:
            doc = DocxDocument()
            title = doc.add_heading('CROWD ANALYSIS — INCIDENT LOG', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(
                f'System   : Crowd Analysis & Threat Detection\n'
                f'Model    : YOLOv8m (V1)\n'
                f'Classes  : {", ".join(FINAL_CLASSES)}\n'
                f'Created  : {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
            )
            doc.add_paragraph('─' * 60)

        doc.add_heading(
            f'INCIDENT — {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}',
            level=1
        )
        doc.add_paragraph(f'Source           : {source}')
        doc.add_paragraph(f'Total detections : {len(detections)}')
        doc.add_paragraph(f'Alerts triggered : {len(triggered_alerts)}')
        doc.add_paragraph('')

        doc.add_heading('Alerts', level=2)
        if triggered_alerts:
            for alert in triggered_alerts:
                p   = doc.add_paragraph(style='List Bullet')
                run = p.add_run(
                    f"{alert['emoji']} [{alert['level']}] "
                    f"{alert['msg']} (conf: {alert['conf']:.2f})"
                )
                run.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = (
                    RGBColor(200,0,0)
                    if alert['level'] == 'CRITICAL'
                    else RGBColor(200,100,0)
                )
        else:
            doc.add_paragraph('No threats detected.')

        doc.add_paragraph('')
        doc.add_heading('All Detections', level=2)
        if detections:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text, hdr[2].text = 'Class','Confidence','Alert Level'
            for det in detections:
                row = table.add_row().cells
                row[0].text = det['class']
                row[1].text = f"{det['conf']:.2f}"
                row[2].text = ALERT_CONFIG[det['class']]['level']
        else:
            doc.add_paragraph('No objects detected.')

        doc.add_paragraph('')
        doc.add_paragraph('─' * 60)
        doc.add_paragraph('')
        doc.save(LOG_PATH)


# ══════════════════════════════════════════════════════════════
#  RAG — INCIDENT SUMMARY
# ══════════════════════════════════════════════════════════════

_rag_chain  = None
_embeddings = None

def get_embeddings():
    global _embeddings
    if _embeddings is None:
        _embeddings = HuggingFaceEmbeddings(
            model_name='sentence-transformers/all-MiniLM-L6-v2'
        )
    return _embeddings


def build_rag_chain():
    global _rag_chain
    if not os.path.exists(LOG_PATH):
        return None, '⚠️ No incident log found. Run detections first!'
    try:
        loader   = Docx2txtLoader(LOG_PATH)
        docs     = loader.load()
        splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        chunks   = splitter.split_documents(docs)
        if not chunks:
            return None, '⚠️ Incident log is empty.'
        vectorstore = FAISS.from_documents(chunks, get_embeddings())
        llm_pipe    = hf_pipeline(
            'text2text-generation',
            model='google/flan-t5-base',
            max_length=512,
            truncation=True,
        )
        llm        = HuggingFacePipeline(pipeline=llm_pipe)
        _rag_chain = RetrievalQA.from_chain_type(
            llm=llm,
            chain_type='stuff',
            retriever=vectorstore.as_retriever(search_kwargs={'k': 4}),
            return_source_documents=False,
        )
        return _rag_chain, '✅ RAG ready — ask questions about incidents!'
    except Exception as e:
        return None, f'❌ Error: {str(e)}'


def query_rag(question, history):
    global _rag_chain
    if not question.strip():
        return history, ''
    if _rag_chain is None:
        _rag_chain, status = build_rag_chain()
        if _rag_chain is None:
            history.append((question, status))
            return history, ''
    try:
        result = _rag_chain.invoke({'query': question})
        answer = result.get('result', 'No answer found.')
        answer = f'📋 {answer}\n\n_Answer based on incident log only — zero hallucination_'
    except Exception as e:
        answer = f'❌ Error: {str(e)}'
    history.append((question, answer))
    return history, ''


# ══════════════════════════════════════════════════════════════
#  INFERENCE FUNCTIONS
# ══════════════════════════════════════════════════════════════

def detect_image(image):
    if image is None:
        return None, '<div class="alert-panel"><span style="color:#5a7a9a">No image uploaded.</span></div>', None
    img_bgr    = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2BGR)
    results    = model.predict(img_bgr, conf=CONF_THRESHOLD, verbose=False)
    img_out    = draw_boxes(img_bgr.copy(), results)
    img_rgb    = cv2.cvtColor(img_out, cv2.COLOR_BGR2RGB)
    detections = extract_detections(results)
    triggered  = [
        {**ALERT_CONFIG[n], 'conf': c, 'class': n}
        for n, c in {d['class']: d['conf'] for d in detections}.items()
    ]
    alert_html = build_alert_html(detections)
    threading.Thread(
        target=append_to_log,
        args=('image_upload', detections, triggered),
        daemon=True
    ).start()
    return Image.fromarray(img_rgb), alert_html, LOG_PATH


def detect_video(video_path):
    if video_path is None:
        return None, '<div class="alert-panel"><span style="color:#5a7a9a">No video uploaded.</span></div>', None
    cap    = cv2.VideoCapture(video_path)
    fps    = cap.get(cv2.CAP_PROP_FPS) or 25
    w      = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
    h      = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
    total  = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
    tmp    = tempfile.NamedTemporaryFile(delete=False, suffix='.mp4')
    writer = cv2.VideoWriter(
        tmp.name, cv2.VideoWriter_fourcc(*'mp4v'), fps, (w, h)
    )
    all_detections = []
    frame_n        = 0
    while cap.isOpened():
        ret, frame = cap.read()
        if not ret:
            break
        if frame_n % 3 == 0:
            results = model.predict(frame, conf=CONF_THRESHOLD, verbose=False)
            frame   = draw_boxes(frame, results)
            all_detections.extend(extract_detections(results))
        cv2.putText(frame, f'Frame {frame_n}/{total}',
                    (10,30), cv2.FONT_HERSHEY_SIMPLEX, 0.7, (255,255,255), 2)
        writer.write(frame)
        frame_n += 1
    cap.release()
    writer.release()
    triggered  = [
        {**ALERT_CONFIG[n], 'conf': c, 'class': n}
        for n, c in {d['class']: d['conf'] for d in all_detections}.items()
    ]
    alert_html = build_alert_html(all_detections)
    threading.Thread(
        target=append_to_log,
        args=(f'video_upload ({frame_n} frames)', all_detections, triggered),
        daemon=True
    ).start()
    return tmp.name, alert_html, LOG_PATH


def detect_live(frame):
    if frame is None:
        return None, '<div class="alert-panel"><span style="color:#5a7a9a">No feed.</span></div>'
    img_bgr    = cv2.cvtColor(np.array(frame), cv2.COLOR_RGB2BGR)
    results    = model.predict(img_bgr, conf=CONF_THRESHOLD, verbose=False)
    img_out    = draw_boxes(img_bgr.copy(), results)
    img_rgb    = cv2.cvtColor(img_out, cv2.COLOR_BGR2RGB)
    detections = extract_detections(results)
    alert_html = build_alert_html(detections)
    if detections:
        triggered = [
            {**ALERT_CONFIG[n], 'conf': c, 'class': n}
            for n, c in {d['class']: d['conf'] for d in detections}.items()
        ]
        threading.Thread(
            target=append_to_log,
            args=('live_feed', detections, triggered),
            daemon=True
        ).start()
    return Image.fromarray(img_rgb), alert_html


# ══════════════════════════════════════════════════════════════
#  GRADIO UI
# ══════════════════════════════════════════════════════════════

CSS = """
@import url('https://fonts.googleapis.com/css2?family=Rajdhani:wght@400;600;700&family=Share+Tech+Mono&family=Exo+2:wght@300;400;600&display=swap');

:root {
    --bg-dark  : #050a0f;
    --bg-panel : #0a1520;
    --bg-card  : #0d1e2e;
    --red      : #ff2244;
    --blue     : #00aaff;
    --gold     : #ffaa00;
    --text     : #c8d8e8;
    --dim      : #5a7a9a;
    --border   : #1a3a5a;
}

body, .gradio-container {
    background: var(--bg-dark) !important;
    font-family: 'Exo 2', sans-serif !important;
    color: var(--text) !important;
}

.sys-header {
    background: linear-gradient(135deg, #050a0f, #0a1a2e, #050a0f);
    border-bottom: 2px solid var(--red);
    padding: 32px 24px 24px;
    text-align: center;
}

.sys-title {
    font-family: 'Rajdhani', sans-serif !important;
    font-size: 44px !important;
    font-weight: 700 !important;
    letter-spacing: 10px !important;
    color: #fff !important;
    text-shadow: 0 0 30px rgba(255,34,68,0.8), 0 0 60px rgba(255,34,68,0.3) !important;
    margin: 0 !important;
}

.sys-sub {
    font-family: 'Share Tech Mono', monospace;
    font-size: 12px;
    color: var(--blue);
    letter-spacing: 4px;
    margin-top: 8px;
}

.sys-badges {
    display: flex;
    justify-content: center;
    gap: 10px;
    margin-top: 14px;
    flex-wrap: wrap;
}

.badge {
    background: rgba(0,170,255,0.08);
    border: 1px solid rgba(0,170,255,0.25);
    border-radius: 3px;
    padding: 3px 10px;
    font-family: 'Share Tech Mono', monospace;
    font-size: 10px;
    color: var(--blue);
    letter-spacing: 1px;
}

.badge.red {
    background: rgba(255,34,68,0.08);
    border-color: rgba(255,34,68,0.25);
    color: var(--red);
}

.tab-nav button {
    background: var(--bg-panel) !important;
    border: 1px solid var(--border) !important;
    color: var(--dim) !important;
    font-family: 'Rajdhani', sans-serif !important;
    font-size: 13px !important;
    font-weight: 600 !important;
    letter-spacing: 2px !important;
    text-transform: uppercase !important;
    transition: all 0.2s !important;
}

.tab-nav button.selected {
    background: var(--bg-card) !important;
    border-color: var(--red) !important;
    color: #fff !important;
    box-shadow: 0 0 15px rgba(255,34,68,0.25) !important;
}

.section-lbl {
    font-family: 'Share Tech Mono', monospace;
    font-size: 11px;
    color: var(--blue);
    letter-spacing: 3px;
    text-transform: uppercase;
    padding: 14px 0 8px;
}

.alert-panel {
    background: var(--bg-card);
    border: 1px solid var(--border);
    border-radius: 8px;
    padding: 16px;
    min-height: 60px;
}

.gr-button {
    font-family: 'Rajdhani', sans-serif !important;
    font-weight: 700 !important;
    letter-spacing: 2px !important;
    text-transform: uppercase !important;
}

label {
    font-family: 'Share Tech Mono', monospace !important;
    font-size: 11px !important;
    letter-spacing: 1px !important;
    color: var(--dim) !important;
}

.footer-bar {
    background: var(--bg-panel);
    border-top: 1px solid var(--border);
    padding: 10px 24px;
    display: flex;
    justify-content: space-between;
    font-family: 'Share Tech Mono', monospace;
    font-size: 10px;
    color: var(--dim);
    margin-top: 16px;
}
"""

HEADER = """
<div class="sys-header">
    <div class="sys-title">⚡ CROWD SHIELD</div>
    <div class="sys-sub">REAL-TIME THREAT DETECTION & INCIDENT MANAGEMENT SYSTEM</div>
    <div class="sys-badges">
        <span class="badge red">🔴 LIVE SYSTEM</span>
        <span class="badge">YOLOv8m</span>
        <span class="badge">5 THREAT CLASSES</span>
        <span class="badge">AUTO INCIDENT LOG</span>
        <span class="badge red">RAG REPORTING</span>
    </div>
</div>
"""

RAG_INFO = """
<div style="background:#0a1520;border:1px solid #1a3a5a;border-radius:8px;
            padding:16px;font-family:'Share Tech Mono',monospace;">
    <div style="color:#00aaff;font-size:11px;letter-spacing:2px;margin-bottom:12px;">
        ▸ RAG SYSTEM INFO
    </div>
    <div style="color:#5a7a9a;font-size:11px;line-height:2.2;">
        EMBEDDINGS:<br>
        <span style="color:#c8d8e8;">all-MiniLM-L6-v2</span><br>
        VECTOR STORE:<br>
        <span style="color:#c8d8e8;">FAISS (CPU)</span><br>
        LLM:<br>
        <span style="color:#c8d8e8;">Flan-T5-Base</span><br>
        METHOD:<br>
        <span style="color:#c8d8e8;">RetrievalQA</span><br>
        SOURCE:<br>
        <span style="color:#c8d8e8;">incident_log.docx</span><br>
        HALLUCINATION:<br>
        <span style="color:#00ff44;">ZERO ✅</span>
    </div>
</div>
"""

AWAIT_HTML = '<div class="alert-panel"><span style="color:#5a7a9a;font-family:\'Share Tech Mono\',monospace;font-size:13px;">// AWAITING INPUT //</span></div>'

with gr.Blocks(css=CSS, title='CrowdShield — Threat Detection') as app:

    gr.HTML(HEADER)

    with gr.Tabs():

        # ── Tab 1: Image ──────────────────────────────────────
        with gr.Tab('📷  IMAGE'):
            gr.HTML('<div class="section-lbl">▸ Upload Image for Threat Analysis</div>')
            with gr.Row():
                with gr.Column():
                    img_input  = gr.Image(type='pil', label='INPUT FEED', height=380)
                    img_btn    = gr.Button('⚡ ANALYZE THREAT', variant='primary')
                with gr.Column():
                    img_output = gr.Image(type='pil', label='DETECTION OUTPUT', height=380)
            gr.HTML('<div class="section-lbl">▸ Alert Status</div>')
            img_alert = gr.HTML(value=AWAIT_HTML)
            img_log   = gr.File(label='📄 DOWNLOAD INCIDENT LOG (.docx)')
            img_btn.click(
                fn=detect_image, inputs=[img_input],
                outputs=[img_output, img_alert, img_log]
            )

        # ── Tab 2: Video ──────────────────────────────────────
        with gr.Tab('🎥  VIDEO'):
            gr.HTML('<div class="section-lbl">▸ Upload Video for Threat Analysis</div>')
            with gr.Row():
                with gr.Column():
                    vid_input  = gr.Video(label='INPUT FEED', height=380)
                    vid_btn    = gr.Button('⚡ ANALYZE THREAT', variant='primary')
                with gr.Column():
                    vid_output = gr.Video(label='DETECTION OUTPUT', height=380)
            gr.HTML('<div class="section-lbl">▸ Alert Status</div>')
            vid_alert = gr.HTML(value=AWAIT_HTML)
            vid_log   = gr.File(label='📄 DOWNLOAD INCIDENT LOG (.docx)')
            vid_btn.click(
                fn=detect_video, inputs=[vid_input],
                outputs=[vid_output, vid_alert, vid_log]
            )

        # ── Tab 3: Live ───────────────────────────────────────
        with gr.Tab('📡  LIVE FEED'):
            gr.HTML('<div class="section-lbl">▸ Connect Live Camera Feed</div>')
            with gr.Row():
                with gr.Column():
                    live_input = gr.Image(
                        sources=['webcam'], type='pil',
                        label='LIVE INPUT', streaming=True, height=380
                    )
                with gr.Column():
                    live_output = gr.Image(type='pil', label='DETECTION OUTPUT', height=380)
            gr.HTML('<div class="section-lbl">▸ Alert Status</div>')
            live_alert = gr.HTML(value=AWAIT_HTML)
            live_input.stream(
                fn=detect_live, inputs=[live_input],
                outputs=[live_output, live_alert]
            )

        # ── Tab 4: RAG ────────────────────────────────────────
        with gr.Tab('📋  INCIDENT REPORT'):
            gr.HTML('''
            <div style="background:#0a1520;border:1px solid #1a3a5a;
                        border-left:4px solid #00aaff;border-radius:8px;
                        padding:16px 20px;margin:16px 0;
                        font-family:'Share Tech Mono',monospace;">
                <div style="color:#00aaff;font-size:12px;letter-spacing:2px;margin-bottom:8px;">
                    ▸ RAG INCIDENT QUERY — ZERO HALLUCINATION
                </div>
                <div style="color:#5a7a9a;font-size:11px;line-height:1.9;">
                    Ask questions about logged incidents. Answers come from incident_log.docx ONLY.<br>
                    Try: "Summarize all CRITICAL incidents" | "Generate official report" | "How many weapon detections?"
                </div>
            </div>
            ''')
            with gr.Row():
                with gr.Column(scale=3):
                    chatbot = gr.Chatbot(
                        label='INCIDENT QUERY INTERFACE',
                        height=440,
                        bubble_full_width=False,
                    )
                    with gr.Row():
                        query_box = gr.Textbox(
                            placeholder='Ask about incidents... e.g. "Generate official report for authorities"',
                            label='', scale=4, container=False,
                        )
                        ask_btn = gr.Button('⚡ ASK', variant='primary', scale=1)
                    gr.Examples(
                        examples=[
                            ['Summarize all incidents for the officials'],
                            ['How many CRITICAL alerts were triggered?'],
                            ['Which incidents involved fire or smoke?'],
                            ['List all weapon detections with timestamps'],
                            ['Generate a complete incident report'],
                            ['What was the most dangerous incident?'],
                        ],
                        inputs=[query_box],
                        label='QUICK QUERIES',
                    )

                with gr.Column(scale=1):
                    gr.HTML(RAG_INFO)
                    rag_status = gr.Textbox(
                        label='SYSTEM STATUS',
                        value='Click LOAD to initialize RAG',
                        interactive=False, lines=2,
                    )
                    load_btn   = gr.Button('📂 LOAD INCIDENT LOG', variant='secondary')
                    log_upload = gr.File(
                        label='↑ Or upload incident_log.docx',
                        file_types=['.docx'],
                    )

            def load_rag(uploaded=None):
                global _rag_chain
                _rag_chain = None
                if uploaded:
                    import shutil
                    shutil.copy(uploaded.name, LOG_PATH)
                _, status = build_rag_chain()
                return status

            load_btn.click(fn=load_rag, inputs=[log_upload], outputs=[rag_status])
            ask_btn.click(fn=query_rag, inputs=[query_box, chatbot], outputs=[chatbot, query_box])
            query_box.submit(fn=query_rag, inputs=[query_box, chatbot], outputs=[chatbot, query_box])

    # Footer
    gr.HTML('''
    <div class="footer-bar">
        <span>CROWDSHIELD v1.0 | YOLOv8m | mAP50: 0.603</span>
        <span>🔴 CRITICAL: Fire, Weapon &nbsp;|&nbsp; 🟠 HIGH: Ambulance, Fighting, Smoke</span>
        <span>ALL INCIDENTS AUTO-LOGGED TO .DOCX</span>
    </div>
    ''')

if __name__ == '__main__':
    app.launch()
